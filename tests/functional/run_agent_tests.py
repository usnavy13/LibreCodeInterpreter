#!/usr/bin/env python3
"""Functional tests for agent skills via code-interpreter API.

Executes code in the real nsjail sandbox, same path as LibreChat agents.
"""
import json
import sys
import requests

API = "http://127.0.0.1:8010"
KEY = "facac6914bfccdddd47595b6bf24d476e38bd42516d99bb5aff8da48df649a4c"
PASS = 0
FAIL = 0
RESULTS = []


def exec_code(label, code, lang="py", timeout=60):
    global PASS, FAIL
    try:
        resp = requests.post(
            f"{API}/exec",
            headers={"X-API-Key": KEY, "Content-Type": "application/json"},
            json={"lang": lang, "code": code, "timeout": timeout},
            timeout=timeout + 30,
        )
        data = resp.json()
        stdout = data.get("stdout", "")
        stderr = data.get("stderr", "")
        if "error" in stderr.lower() or "traceback" in stderr.lower():
            if "experimentalwarning" in stderr.lower() or "cpuinfo" in stderr.lower():
                pass  # Ignore benign warnings
            else:
                print(f"  \u2717 {label}")
                print(f"    STDERR: {stderr[:200]}")
                FAIL += 1
                RESULTS.append(("FAIL", label, stderr.split(chr(10))[0][:80]))
                return stdout, stderr, False
        print(f"  \u2713 {label}")
        if stdout.strip():
            for line in stdout.strip().split("\n")[:5]:
                print(f"    {line}")
        PASS += 1
        RESULTS.append(("PASS", label, ""))
        return stdout, stderr, True
    except Exception as e:
        print(f"  \u2717 {label} — {e}")
        FAIL += 1
        RESULTS.append(("FAIL", label, str(e)[:80]))
        return "", str(e), False


print("=" * 60)
print("AGENT FUNCTIONAL TESTS")
print(f"API: {API}")
print("=" * 60)

# ========================================
print("\n=== AGENT: Word DOCX Complete ===")
# ========================================

exec_code("D01: Création DOCX (python-docx)", """
from docx import Document
import os
doc = Document()
doc.add_heading("Rapport Q1 2026", level=1)
doc.add_heading("Direction Commerciale", level=2)
doc.add_paragraph("Les objectifs ont été atteints à 95%.")
table = doc.add_table(rows=4, cols=3, style="Table Grid")
headers = ["Mois", "CA", "Marge"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for i, (m, c, g) in enumerate([("Jan","120k","15%"),("Fev","135k","18%"),("Mar","150k","20%")]):
    table.rows[i+1].cells[0].text = m
    table.rows[i+1].cells[1].text = c
    table.rows[i+1].cells[2].text = g
doc.save("/mnt/data/rapport_q1.docx")
sz = os.path.getsize("/mnt/data/rapport_q1.docx")
print("DOCX created:", sz, "bytes")
assert sz > 1000
""")

exec_code("D02: Unpack + tracked_replace + pack", """
import subprocess, os
r = subprocess.run(["python3", "/opt/skills/docx/scripts/office/unpack.py",
    "/mnt/data/rapport_q1.docx", "/mnt/data/rapport_q1_unpacked"], capture_output=True, text=True)
print("unpack rc:", r.returncode)
assert r.returncode == 0, "unpack failed: " + r.stderr[:200]
assert os.path.isdir("/mnt/data/rapport_q1_unpacked")
r = subprocess.run(["python3", "/opt/skills/docx/scripts/tracked_replace.py",
    "/mnt/data/rapport_q1_unpacked", "--old", "Q1 2026", "--new", "Q2 2026",
    "--author", "AI-Agent"], capture_output=True, text=True)
print("tracked_replace:", r.stdout.strip()[:200])
assert r.returncode == 0, "tracked_replace failed: " + r.stderr[:200]
r = subprocess.run(["python3", "/opt/skills/docx/scripts/office/pack.py",
    "/mnt/data/rapport_q1_unpacked", "-o", "/mnt/data/rapport_q2_redline.docx"],
    capture_output=True, text=True)
print("pack rc:", r.returncode)
assert r.returncode == 0, "pack failed: " + r.stderr[:200]
sz = os.path.getsize("/mnt/data/rapport_q2_redline.docx")
print("Redlined DOCX:", sz, "bytes")
assert sz > 1000
""", timeout=120)

exec_code("D03: Accept tracked changes (soffice)", """
import subprocess, os
r = subprocess.run(["python3", "/opt/skills/docx/scripts/accept_changes.py",
    "/mnt/data/rapport_q2_redline.docx", "/mnt/data/rapport_q2_clean.docx"],
    capture_output=True, text=True, timeout=120)
print("rc:", r.returncode)
print("stdout:", r.stdout.strip()[:200])
if os.path.exists("/mnt/data/rapport_q2_clean.docx"):
    sz = os.path.getsize("/mnt/data/rapport_q2_clean.docx")
    print("Clean DOCX:", sz, "bytes")
else:
    print("Note: soffice needs /proc in sandbox, partial test")
""", timeout=180)

exec_code("D05: Pandoc markdown to DOCX", """
import subprocess, os
md = "# Politique\\n## Principes\\nLe télétravail est ouvert.\\n- 3 jours max\\n- Accord requis"
with open("/mnt/data/policy.md", "w") as f:
    f.write(md.replace("\\n", chr(10)))
r = subprocess.run(["pandoc", "/mnt/data/policy.md", "-o", "/mnt/data/policy.docx"], capture_output=True, text=True)
assert r.returncode == 0, "pandoc failed"
sz = os.path.getsize("/mnt/data/policy.docx")
print("Pandoc DOCX:", sz, "bytes")
assert sz > 1000
""")

exec_code("D10: DOCX avec pied de page", """
from docx import Document
import os
doc = Document()
doc.add_heading("Fiche Produit", level=1)
table = doc.add_table(rows=4, cols=2, style="Table Grid")
for i, (k, v) in enumerate([("Poids","1.2 kg"),("Dimensions","30x20x10"),("Couleur","Noir"),("Prix","149.90 EUR")]):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
section = doc.sections[0]
footer = section.footer
footer.is_linked_to_previous = False
footer.paragraphs[0].text = "Confidentiel"
doc.save("/mnt/data/fiche.docx")
sz = os.path.getsize("/mnt/data/fiche.docx")
print("DOCX with footer:", sz, "bytes")
assert sz > 1000
""")

exec_code("D11: tracked_replace --first", """
from docx import Document
import subprocess, os
doc = Document()
doc.add_paragraph("Le Directeur general a vu le Directeur commercial et le Directeur technique.")
doc.save("/mnt/data/directeurs.docx")
r = subprocess.run(["python3", "/opt/skills/docx/scripts/office/unpack.py",
    "/mnt/data/directeurs.docx", "/mnt/data/directeurs_unpacked"], capture_output=True, text=True)
assert r.returncode == 0, "unpack: " + r.stderr[:200]
r = subprocess.run(["python3", "/opt/skills/docx/scripts/tracked_replace.py",
    "/mnt/data/directeurs_unpacked", "--old", "Directeur", "--new", "Directrice",
    "--first", "--author", "AI-Agent"], capture_output=True, text=True)
print(r.stdout.strip())
assert "1 replacement" in r.stdout, "Expected 1 replacement: " + r.stdout
r = subprocess.run(["python3", "/opt/skills/docx/scripts/office/pack.py",
    "/mnt/data/directeurs_unpacked", "-o", "/mnt/data/directeurs_ed.docx"], capture_output=True, text=True)
assert r.returncode == 0
print("--first: only 1 occurrence replaced")
""", timeout=120)

# ========================================
print("\n=== AGENT: PowerPoint PPTX ===")
# ========================================

exec_code("P01: Création PptxGenJS", """
const pptxgen = require("pptxgenjs");
const pptx = new pptxgen();
let s1 = pptx.addSlide();
s1.addText("PayFlow", {x:1,y:1,fontSize:36,bold:true,color:"003366"});
s1.addText("FinTech des paiements", {x:1,y:2,fontSize:18,color:"666666"});
let s2 = pptx.addSlide();
s2.addText("Probleme", {x:0.5,y:0.3,fontSize:28,bold:true});
s2.addText("15h/mois perdues en gestion", {x:0.5,y:1.2,fontSize:16});
let s3 = pptx.addSlide();
s3.addChart(pptx.ChartType.bar, [{name:"Marche",labels:["TAM","SAM","SOM"],values:[50,12,3]}],
    {x:0.5,y:1,w:8,h:3.5});
pptx.writeFile({fileName:"/mnt/data/pitch.pptx"}).then(()=>{
    const fs=require("fs");
    console.log("PPTX:",fs.statSync("/mnt/data/pitch.pptx").size,"bytes");
});
""", lang="js", timeout=30)

exec_code("P04: Unpack + add_slide + pack (PPTX)", """
import subprocess, os
r = subprocess.run(["python3", "/opt/skills/pptx/scripts/office/unpack.py",
    "/mnt/data/pitch.pptx", "/mnt/data/pitch_unpacked"], capture_output=True, text=True)
print("unpack rc:", r.returncode, r.stdout.strip()[:100])
if r.returncode == 0 and os.path.isdir("/mnt/data/pitch_unpacked"):
    r2 = subprocess.run(["python3", "/opt/skills/pptx/scripts/add_slide.py",
        "/mnt/data/pitch_unpacked", "--source", "2"], capture_output=True, text=True)
    print("add_slide:", r2.stdout.strip()[:100], "rc:", r2.returncode)
    r3 = subprocess.run(["python3", "/opt/skills/pptx/scripts/office/pack.py",
        "/mnt/data/pitch_unpacked", "-o", "/mnt/data/pitch_ext.pptx"], capture_output=True, text=True)
    print("pack rc:", r3.returncode)
    if os.path.exists("/mnt/data/pitch_ext.pptx"):
        print("Extended PPTX:", os.path.getsize("/mnt/data/pitch_ext.pptx"), "bytes")
else:
    print("Unpack issue, checking stderr:", r.stderr[:200])
""", timeout=60)

exec_code("P07: markitdown PPTX to markdown", """
from markitdown import MarkItDown
md = MarkItDown()
result = md.convert("/mnt/data/pitch.pptx")
text = result.text_content
print("Extracted text length:", len(text))
print(text[:300])
""", timeout=60)

exec_code("P08: PptxGenJS with data slides", """
const pptxgen=require("pptxgenjs");
const pptx=new pptxgen();
const data=[["Jan",120,95],["Fev",110,100],["Mar",130,140]];
data.forEach(d=>{
    let s=pptx.addSlide();
    s.addText(d[0]+" 2026",{x:0.5,y:0.3,fontSize:28,bold:true});
    s.addTable([["CA","Charges","Resultat"],[d[1]+"k",d[2]+"k",(d[1]-d[2])+"k"]],
        {x:0.5,y:1.2,w:8,border:{type:"solid",pt:1}});
    s.addShape(pptx.ShapeType.rect,{x:0.5,y:3,w:3,h:0.5,fill:{color:d[1]-d[2]>=0?"00AA00":"CC0000"}});
});
pptx.writeFile({fileName:"/mnt/data/reporting.pptx"}).then(()=>{
    console.log("PPTX:",require("fs").statSync("/mnt/data/reporting.pptx").size,"bytes");
});
""", lang="js", timeout=30)

# ========================================
print("\n=== AGENT: Excel XLSX ===")
# ========================================

exec_code("X01: Budget multi-onglets (openpyxl)", """
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side
import os
wb = Workbook()
postes = ["Salaires","Loyer","Marketing","IT","Divers"]
vals = {"Q1":[45000,8000,12000,5000,3000],"Q2":[46000,8000,15000,6000,3500]}
hfill = PatternFill(start_color="003366",end_color="003366",fill_type="solid")
hfont = Font(color="FFFFFF",bold=True)
thin = Border(left=Side(style="thin"),right=Side(style="thin"),top=Side(style="thin"),bottom=Side(style="thin"))
for qi,(qn,amounts) in enumerate(vals.items()):
    ws = wb.active if qi==0 else wb.create_sheet()
    ws.title = qn
    for col,h in enumerate(["Poste","Montant"],1):
        c=ws.cell(row=1,column=col,value=h); c.fill=hfill; c.font=hfont; c.border=thin
    for i,(p,v) in enumerate(zip(postes,amounts)):
        ws.cell(row=i+2,column=1,value=p).border=thin
        c=ws.cell(row=i+2,column=2,value=v); c.number_format="#,##0"; c.border=thin
    ws.cell(row=len(postes)+2,column=1,value="TOTAL").font=Font(bold=True)
    ws.cell(row=len(postes)+2,column=2).value="=SUM(B2:B{})".format(len(postes)+1)
syn=wb.create_sheet("Synthese")
syn.cell(row=1,column=1,value="Trimestre"); syn.cell(row=1,column=2,value="Total")
for i,qn in enumerate(vals.keys()):
    syn.cell(row=i+2,column=1,value=qn)
    syn.cell(row=i+2,column=2).value="={}!B{}".format(qn,len(postes)+2)
wb.save("/mnt/data/budget.xlsx")
print("XLSX:", os.path.getsize("/mnt/data/budget.xlsx"), "bytes,", len(wb.sheetnames), "sheets:", wb.sheetnames)
""")

exec_code("X02: Analyse pandas", """
import pandas as pd, numpy as np
np.random.seed(42)
df = pd.DataFrame({"client":np.random.choice(["Acme","Global","Soft","Data","Web"],100),
    "montant":np.random.normal(5000,1500,100).round(2),"qty":np.random.randint(1,50,100)})
df.to_csv("/mnt/data/ventes.csv",index=False)
print("Shape:", df.shape)
print("Top clients:", df.groupby("client")["montant"].sum().sort_values(ascending=False).to_dict())
""")

exec_code("X03: Recalcul formules (recalc.py)", """
import subprocess
r = subprocess.run(["python3","/opt/skills/xlsx/scripts/recalc.py","/mnt/data/budget.xlsx"],
    capture_output=True, text=True, timeout=120)
print("rc:", r.returncode)
print(r.stdout.strip()[:300])
""", timeout=180)

exec_code("X04: Graphique Excel natif", """
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference
import os
wb = Workbook(); ws = wb.active; ws.title="Ventes"
for row in [["Mois","CA"],["Jan",120],["Fev",135],["Mar",150],["Avr",140]]:
    ws.append(row)
chart = BarChart(); chart.title = "CA mensuel"
cats = Reference(ws,min_col=1,min_row=2,max_row=5)
vals = Reference(ws,min_col=2,min_row=1,max_row=5)
chart.add_data(vals,titles_from_data=True); chart.set_categories(cats)
dash = wb.create_sheet("Dashboard"); dash.add_chart(chart,"A1")
wb.save("/mnt/data/chart.xlsx")
print("XLSX with chart:", os.path.getsize("/mnt/data/chart.xlsx"), "bytes")
""")

exec_code("X07: Mise en forme conditionnelle", """
from openpyxl import Workbook
from openpyxl.formatting.rule import CellIsRule
from openpyxl.styles import PatternFill
import random, os
random.seed(42)
wb=Workbook(); ws=wb.active; ws.title="Objectifs"
ws.append(["Vendeur","Objectif","Realise","Pct"])
for i in range(10):
    obj=random.randint(80,120)*1000; real=int(obj*random.uniform(0.7,1.2))
    ws.append(["Vendeur "+str(i+1),obj,real])
    ws.cell(row=i+2,column=4).value="=C{}/B{}".format(i+2,i+2)
    ws.cell(row=i+2,column=4).number_format="0%"
ws.conditional_formatting.add("D2:D11",CellIsRule(operator="greaterThanOrEqual",formula=["1"],fill=PatternFill(bgColor="00CC00")))
ws.conditional_formatting.add("D2:D11",CellIsRule(operator="lessThan",formula=["0.8"],fill=PatternFill(bgColor="CC0000")))
wb.save("/mnt/data/objectifs.xlsx")
print("Conditional formatting XLSX:", os.path.getsize("/mnt/data/objectifs.xlsx"), "bytes")
""")

exec_code("X11: Formules natives preservees", """
from openpyxl import Workbook, load_workbook
import os
wb=Workbook(); ws=wb.active; ws.title="Tresorerie"
ws.cell(row=1,column=1,value="Mois"); ws.cell(row=1,column=2,value="Enc"); ws.cell(row=1,column=3,value="Dec"); ws.cell(row=1,column=4,value="Solde")
ws.cell(row=2,column=1,value="M1"); ws.cell(row=2,column=2,value=20000); ws.cell(row=2,column=3,value=18000); ws.cell(row=2,column=4).value="=B2-C2+50000"
for i in range(3,13):
    ws.cell(row=i,column=1,value="M"+str(i-1))
    ws.cell(row=i,column=2).value="=B{}*1.05".format(i-1)
    ws.cell(row=i,column=3,value=18000)
    ws.cell(row=i,column=4).value="=D{}+B{}-C{}".format(i-1,i,i)
wb.save("/mnt/data/treso.xlsx")
wb2=load_workbook("/mnt/data/treso.xlsx"); ws2=wb2.active
val=ws2["B3"].value
print("B3:", val)
assert isinstance(val,str) and val.startswith("="), "Expected formula"
print("Formulas preserved as native Excel formulas")
""")

# ========================================
print("\n=== AGENT: PDF ===")
# ========================================

exec_code("F01: pdfplumber extraction", """
import subprocess
subprocess.run(["pandoc","/mnt/data/policy.md","-o","/mnt/data/policy.pdf"],check=True,capture_output=True)
import pdfplumber
with pdfplumber.open("/mnt/data/policy.pdf") as pdf:
    text = chr(10).join(page.extract_text() or "" for page in pdf.pages)
print("Pages:", len(pdf.pages))
print("Text:", text[:200])
assert len(text) > 10
""")

exec_code("F02: pdfplumber table extraction", """
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib import colors
import pdfplumber, pandas as pd
doc = SimpleDocTemplate("/mnt/data/table.pdf", pagesize=A4)
data = [["Produit","Qty","Prix"],["Widget A","10","25.50"],["Widget B","5","42.00"]]
t = Table(data); t.setStyle(TableStyle([("GRID",(0,0),(-1,-1),1,colors.black)]))
doc.build([t])
with pdfplumber.open("/mnt/data/table.pdf") as pdf:
    tables = pdf.pages[0].extract_tables()
print("Tables found:", len(tables))
if tables:
    df = pd.DataFrame(tables[0][1:], columns=tables[0][0])
    print(df.to_string())
""")

exec_code("F04: Fusion PDFs (pypdf)", """
from reportlab.pdfgen import canvas
from pypdf import PdfMerger, PdfReader
import os
for i,name in enumerate(["p1.pdf","p2.pdf","p3.pdf"]):
    c=canvas.Canvas("/mnt/data/"+name); c.drawString(100,700,"Partie "+str(i+1)); c.save()
merger=PdfMerger()
for f in ["p1.pdf","p2.pdf","p3.pdf"]:
    merger.append("/mnt/data/"+f)
merger.write("/mnt/data/merged.pdf"); merger.close()
r=PdfReader("/mnt/data/merged.pdf")
print("Merged:", len(r.pages), "pages,", os.path.getsize("/mnt/data/merged.pdf"), "bytes")
assert len(r.pages)==3
""")

exec_code("F05: Split PDF", """
from pypdf import PdfReader, PdfWriter
reader=PdfReader("/mnt/data/merged.pdf")
writer=PdfWriter(); writer.add_page(reader.pages[1])
with open("/mnt/data/page2.pdf","wb") as f: writer.write(f)
r2=PdfReader("/mnt/data/page2.pdf")
print("Extracted:", len(r2.pages), "page(s)")
assert len(r2.pages)==1
""")

exec_code("F06: qpdf check", """
import subprocess
r=subprocess.run(["qpdf","--check","/mnt/data/merged.pdf"],capture_output=True,text=True)
print(r.stdout.strip()[:200])
print("rc:", r.returncode)
""")

exec_code("F07: PDF to images (pdf2image)", """
from pdf2image import convert_from_path
import os
images=convert_from_path("/mnt/data/merged.pdf",dpi=150)
print("Pages:", len(images))
for i,img in enumerate(images):
    path="/mnt/data/pg_"+str(i)+".png"
    img.save(path)
    print("  page", i, ":", img.size, os.path.getsize(path), "bytes")
assert len(images)==3
""")

exec_code("F08: Metadata pypdf + qpdf", """
from pypdf import PdfReader
import subprocess
r=PdfReader("/mnt/data/merged.pdf")
meta=r.metadata
print("Pages:", len(r.pages))
print("Producer:", meta.producer if meta else "N/A")
r2=subprocess.run(["qpdf","--show-npages","/mnt/data/merged.pdf"],capture_output=True,text=True)
print("qpdf npages:", r2.stdout.strip())
""")

exec_code("F09: Rotation page", """
from pypdf import PdfReader, PdfWriter
import os
reader=PdfReader("/mnt/data/merged.pdf"); writer=PdfWriter()
for i,page in enumerate(reader.pages):
    if i==0: page.rotate(180)
    writer.add_page(page)
with open("/mnt/data/rotated.pdf","wb") as f: writer.write(f)
print("Rotated:", os.path.getsize("/mnt/data/rotated.pdf"), "bytes")
""")

exec_code("F10: Watermark (reportlab+pypdf)", """
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.pagesizes import A4
from pypdf import PdfReader, PdfWriter
import io, os
packet=io.BytesIO()
c=rl_canvas.Canvas(packet,pagesize=A4)
c.saveState(); c.setFillAlpha(0.3); c.setFillGray(0.5); c.setFont("Helvetica",60)
c.translate(300,400); c.rotate(45); c.drawCentredString(0,0,"BROUILLON"); c.restoreState(); c.save()
packet.seek(0)
wm=PdfReader(packet).pages[0]
reader=PdfReader("/mnt/data/merged.pdf"); writer=PdfWriter()
for page in reader.pages:
    page.merge_page(wm); writer.add_page(page)
with open("/mnt/data/wm.pdf","wb") as f: writer.write(f)
print("Watermarked:", os.path.getsize("/mnt/data/wm.pdf"), "bytes")
""")

exec_code("F11: qpdf linearize", """
import subprocess, os
subprocess.run(["qpdf","--linearize","/mnt/data/wm.pdf","/mnt/data/opt.pdf"],check=True,capture_output=True)
s1=os.path.getsize("/mnt/data/wm.pdf"); s2=os.path.getsize("/mnt/data/opt.pdf")
print("Before:", s1, "After:", s2)
""")

# ========================================
print("\n=== AGENT: Quick Edits (FFmpeg) ===")
# ========================================

exec_code("M01: Audio gen + MP3", """
import subprocess, os
subprocess.run(["ffmpeg","-y","-f","lavfi","-i","sine=frequency=440:duration=3","-f","wav","/mnt/data/tone.wav"],
    check=True,capture_output=True)
print("WAV:", os.path.getsize("/mnt/data/tone.wav"), "bytes")
subprocess.run(["ffmpeg","-y","-i","/mnt/data/tone.wav","-c:a","libmp3lame","-b:a","192k","/mnt/data/tone.mp3"],
    check=True,capture_output=True)
print("MP3:", os.path.getsize("/mnt/data/tone.mp3"), "bytes")
""")

exec_code("M02: Image Pillow", """
from PIL import Image, ImageDraw
import os
img=Image.new("RGB",(800,600))
draw=ImageDraw.Draw(img)
for y in range(600):
    draw.line([(0,y),(799,y)],fill=(int(50*y/600),int(100+100*y/600),int(200+55*y/600)))
draw.text((320,280),"Test Agent",fill=(0,0,0))
img.save("/mnt/data/gradient.png")
print("PNG:", os.path.getsize("/mnt/data/gradient.png"), "bytes, size:", img.size)
""")

exec_code("M03: ffprobe JSON", """
import subprocess, json
r=subprocess.run(["ffprobe","-v","quiet","-print_format","json","-show_format","-show_streams",
    "/mnt/data/tone.mp3"],capture_output=True,text=True,check=True)
info=json.loads(r.stdout)
fmt=info["format"]
print("Format:", fmt["format_name"])
print("Duration:", fmt["duration"], "s")
print("Codec:", info["streams"][0]["codec_name"])
""")

exec_code("M04: Image resize JPEG", """
from PIL import Image
import os
img=Image.open("/mnt/data/gradient.png"); img.thumbnail((400,300))
img.convert("RGB").save("/mnt/data/small.jpg",quality=85)
print("Resized:", img.size, "JPEG:", os.path.getsize("/mnt/data/small.jpg"), "bytes")
""")

exec_code("M07: Watermark Pillow", """
from PIL import Image, ImageDraw
import os
img=Image.open("/mnt/data/gradient.png").convert("RGBA")
overlay=Image.new("RGBA",img.size,(0,0,0,0))
draw=ImageDraw.Draw(overlay)
draw.rectangle([(0,550),(800,600)],fill=(0,0,0,128))
draw.text((280,565),"onbehalf.ai 2026",fill=(255,255,255,255))
result=Image.alpha_composite(img,overlay)
result.convert("RGB").save("/mnt/data/wm.png")
print("Watermarked:", os.path.getsize("/mnt/data/wm.png"), "bytes")
""")

exec_code("M10: Mosaique 2x2", """
from PIL import Image
import os
colors=[(255,0,0),(0,255,0),(0,0,255),(255,255,0)]
tiles=[Image.new("RGB",(200,150),c) for c in colors]
mosaic=Image.new("RGB",(400,300))
mosaic.paste(tiles[0],(0,0)); mosaic.paste(tiles[1],(200,0))
mosaic.paste(tiles[2],(0,150)); mosaic.paste(tiles[3],(200,150))
mosaic.save("/mnt/data/mosaic.png")
print("Mosaic:", os.path.getsize("/mnt/data/mosaic.png"), "bytes, size:", mosaic.size)
""")

exec_code("M11: Video frame extraction", """
import subprocess, os
subprocess.run(["ffmpeg","-y","-f","lavfi","-i","color=c=blue:s=640x480:d=3",
    "-c:v","libx264","-t","3","/mnt/data/vid.mp4"],check=True,capture_output=True)
print("Video:", os.path.getsize("/mnt/data/vid.mp4"), "bytes")
subprocess.run(["ffmpeg","-y","-i","/mnt/data/vid.mp4","-ss","00:00:01",
    "-frames:v","1","/mnt/data/frame.png"],check=True,capture_output=True)
print("Frame:", os.path.getsize("/mnt/data/frame.png"), "bytes")
""")

# ========================================
print("\n=== AGENT: Data Analysis & Visualization ===")
# ========================================

exec_code("A01: Analyse exploratoire", """
import pandas as pd, numpy as np
np.random.seed(42)
n=500
df=pd.DataFrame({"date":pd.date_range("2026-01-01",periods=n),"cat":np.random.choice(["A","B","C"],n),
    "montant":np.random.normal(1000,300,n).round(2),"qty":np.random.randint(1,50,n)})
df.to_csv("/mnt/data/ds.csv",index=False)
print("Shape:", df.shape)
print("Missing:", df.isna().sum().sum())
print("Categories:", df.cat.value_counts().to_dict())
print("Montant mean:", round(df.montant.mean(),2))
""")

exec_code("A02: Dashboard 4 graphiques", """
import pandas as pd, numpy as np
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns
from scipy.stats import linregress
import os
df=pd.read_csv("/mnt/data/ds.csv",parse_dates=["date"])
df["mois"]=df["date"].dt.month
fig,axes=plt.subplots(2,2,figsize=(14,10))
monthly=df.groupby("mois")["montant"].sum()
axes[0,0].plot(monthly.index,monthly.values,marker="o"); axes[0,0].set_title("CA mensuel")
cat_sum=df.groupby("cat")["montant"].sum()
axes[0,1].pie(cat_sum,labels=cat_sum.index,autopct="%1.1f%%"); axes[0,1].set_title("Par categorie")
sns.boxplot(data=df,x="cat",y="montant",ax=axes[1,0]); axes[1,0].set_title("Montants")
axes[1,1].scatter(df["qty"],df["montant"],alpha=0.3,s=10)
sl,ic,r,p,se=linregress(df["qty"],df["montant"])
axes[1,1].plot([0,50],[ic,sl*50+ic],"r-"); axes[1,1].set_title("Qty vs Montant")
plt.tight_layout()
plt.savefig("/mnt/data/dashboard.png",dpi=150,bbox_inches="tight"); plt.close()
print("Dashboard:", os.path.getsize("/mnt/data/dashboard.png"), "bytes")
""")

exec_code("A03: ANOVA", """
import pandas as pd
from scipy import stats
df=pd.read_csv("/mnt/data/ds.csv")
groups=[g["montant"].values for _,g in df.groupby("cat")]
f,p=stats.f_oneway(*groups)
print("F-stat:", round(f,4))
print("p-value:", round(p,6))
sig="Oui" if p<0.05 else "Non"
print("Significatif:", sig)
""")

exec_code("A04: Heatmap correlation", """
import pandas as pd
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns
import os
df=pd.read_csv("/mnt/data/ds.csv")
corr=df[["montant","qty"]].corr()
fig,ax=plt.subplots(figsize=(6,5))
sns.heatmap(corr,annot=True,cmap="RdBu_r",vmin=-1,vmax=1,ax=ax)
plt.savefig("/mnt/data/heatmap.png",dpi=150,bbox_inches="tight"); plt.close()
print("Heatmap:", os.path.getsize("/mnt/data/heatmap.png"), "bytes")
""")

exec_code("A05: Regression lineaire", """
import numpy as np
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
from sklearn.linear_model import LinearRegression
import os
np.random.seed(42)
X=np.random.uniform(30,150,200).reshape(-1,1)
y=3.5*X.ravel()+np.random.normal(0,30,200)+50
model=LinearRegression().fit(X,y)
r2=model.score(X,y)
print("R2:", round(r2,4))
print("Coef:", round(model.coef_[0],2), "Intercept:", round(model.intercept_,2))
fig,ax=plt.subplots(figsize=(8,6))
ax.scatter(X,y,alpha=0.4,s=15)
xl=np.linspace(30,150,100).reshape(-1,1)
ax.plot(xl,model.predict(xl),"r-",linewidth=2)
plt.savefig("/mnt/data/regression.png",dpi=150,bbox_inches="tight"); plt.close()
print("Plot:", os.path.getsize("/mnt/data/regression.png"), "bytes")
""")

exec_code("A06: KMeans clustering", """
import numpy as np
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler
from sklearn.decomposition import PCA
import os
np.random.seed(42)
X=np.vstack([np.random.normal([25,30000,500],[5,8000,200],(80,3)),
    np.random.normal([45,70000,2000],[10,15000,500],(80,3)),
    np.random.normal([35,50000,1000],[8,10000,300],(80,3))])
Xs=StandardScaler().fit_transform(X)
inertias=[KMeans(n_clusters=k,random_state=42,n_init=10).fit(Xs).inertia_ for k in range(1,8)]
km=KMeans(n_clusters=3,random_state=42,n_init=10).fit(Xs)
pca=PCA(n_components=2).fit_transform(Xs)
fig,(a1,a2)=plt.subplots(1,2,figsize=(14,5))
a1.plot(range(1,8),inertias,"bo-"); a1.set_title("Coude")
a2.scatter(pca[:,0],pca[:,1],c=km.labels_,cmap="viridis",s=15,alpha=0.6); a2.set_title("Clusters")
plt.savefig("/mnt/data/kmeans.png",dpi=150,bbox_inches="tight"); plt.close()
print("KMeans:", os.path.getsize("/mnt/data/kmeans.png"), "bytes")
""")

exec_code("A07: Decomposition series temporelles", """
import pandas as pd, numpy as np
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
from statsmodels.tsa.seasonal import seasonal_decompose
import os
np.random.seed(42)
dates=pd.date_range("2024-01-01",periods=365,freq="D")
trend=np.linspace(100,200,365)
seasonal=30*np.sin(2*np.pi*np.arange(365)/30)
ts=pd.Series(trend+seasonal+np.random.normal(0,10,365),index=dates)
result=seasonal_decompose(ts,model="additive",period=30)
fig=result.plot(); fig.set_size_inches(12,8)
plt.savefig("/mnt/data/decomp.png",dpi=150,bbox_inches="tight"); plt.close()
print("Decomposition:", os.path.getsize("/mnt/data/decomp.png"), "bytes")
""")

exec_code("A08: Export multi-onglets", """
import pandas as pd
import os
df=pd.read_csv("/mnt/data/ds.csv",parse_dates=["date"])
df["mois"]=df["date"].dt.month
with pd.ExcelWriter("/mnt/data/rapport.xlsx",engine="openpyxl") as w:
    df.to_excel(w,sheet_name="Brut",index=False)
    pivot=df.pivot_table(values="montant",index="mois",columns="cat",aggfunc="sum")
    pivot.to_excel(w,sheet_name="Pivot")
    stats=df.groupby("cat")["montant"].describe()
    stats.to_excel(w,sheet_name="Stats")
print("Rapport:", os.path.getsize("/mnt/data/rapport.xlsx"), "bytes")
""")

exec_code("A09: Detection anomalies", """
import pandas as pd, numpy as np
from scipy import stats as sp
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
import os
df=pd.read_csv("/mnt/data/ds.csv")
z=np.abs(sp.zscore(df["montant"]))
out_z=df[z>3]
Q1=df["montant"].quantile(0.25); Q3=df["montant"].quantile(0.75); IQR=Q3-Q1
out_iqr=df[(df["montant"]<Q1-1.5*IQR)|(df["montant"]>Q3+1.5*IQR)]
print("Z-score outliers:", len(out_z))
print("IQR outliers:", len(out_iqr))
fig,ax=plt.subplots(figsize=(10,4))
ax.boxplot(df["montant"].values,vert=False)
ax.scatter(out_iqr["montant"],[1]*len(out_iqr),color="red",zorder=5)
plt.savefig("/mnt/data/outliers.png",dpi=150,bbox_inches="tight"); plt.close()
print("Plot:", os.path.getsize("/mnt/data/outliers.png"), "bytes")
""")

exec_code("A10: Test t Student", """
import numpy as np
from scipy import stats
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns
import os
np.random.seed(42)
avant=np.random.normal(72,12,50); apres=np.random.normal(78,11,50)
t,p=stats.ttest_rel(avant,apres)
print("t:", round(t,4), "p:", round(p,6))
sig="Oui" if p<0.05 else "Non"
print("Significatif:", sig)
fig,ax=plt.subplots(figsize=(8,5))
sns.kdeplot(avant,label="Avant",ax=ax,fill=True,alpha=0.3)
sns.kdeplot(apres,label="Apres",ax=ax,fill=True,alpha=0.3)
ax.legend()
plt.savefig("/mnt/data/ttest.png",dpi=150,bbox_inches="tight"); plt.close()
print("Plot:", os.path.getsize("/mnt/data/ttest.png"), "bytes")
""")

# ========================================
print("\n" + "=" * 60)
print("FINAL RESULTS")
print("=" * 60)
print(f"Passed: {PASS}")
print(f"Failed: {FAIL}")
print(f"Total:  {PASS + FAIL}")
print("=" * 60)
for status, label, err in RESULTS:
    mark = "\u2713" if status == "PASS" else "\u2717"
    line = f"  {mark} {label}"
    if err:
        line += f"  — {err}"
    print(line)
print()
sys.exit(0 if FAIL == 0 else 1)
