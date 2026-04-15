#!/usr/bin/env python3
"""Chained pipeline tests using session_id for file persistence."""
import json
import sys
import uuid
import requests

API = "http://127.0.0.1:8010"
KEY = "facac6914bfccdddd47595b6bf24d476e38bd42516d99bb5aff8da48df649a4c"
PASS = 0
FAIL = 0
RESULTS = []


def exec_code(label, code, lang="py", timeout=120, session_id=None):
    global PASS, FAIL
    payload = {"lang": lang, "code": code, "timeout": timeout}
    if session_id:
        payload["session_id"] = session_id
    try:
        resp = requests.post(
            f"{API}/exec",
            headers={"X-API-Key": KEY, "Content-Type": "application/json"},
            json=payload,
            timeout=timeout + 30,
        )
        data = resp.json()
        stdout = data.get("stdout", "")
        stderr = data.get("stderr", "")
        sid = data.get("session_id", "")
        has_error = False
        for err_kw in ["error", "traceback", "exception"]:
            if err_kw in stderr.lower():
                if "experimentalwarning" not in stderr.lower() and "cpuinfo" not in stderr.lower():
                    has_error = True
        if has_error:
            print(f"  \u2717 {label}")
            for line in stderr.strip().split("\n")[:4]:
                print(f"    {line}")
            FAIL += 1
            RESULTS.append(("FAIL", label, stderr.split("\n")[0][:80]))
            return sid, False
        print(f"  \u2713 {label}")
        for line in stdout.strip().split("\n")[:3]:
            print(f"    {line}")
        PASS += 1
        RESULTS.append(("PASS", label, ""))
        return sid, True
    except Exception as e:
        print(f"  \u2717 {label} -- {e}")
        FAIL += 1
        RESULTS.append(("FAIL", label, str(e)[:80]))
        return None, False


print("=" * 60)
print("CHAINED PIPELINE TESTS (with session_id)")
print("=" * 60)

# ====================================================================
print("\n=== Pipeline DOCX: create -> unpack -> tracked_replace -> pack -> validate -> accept -> PDF ===")
# ====================================================================
sid = str(uuid.uuid4())

sid, ok = exec_code("DOCX-1: Create DOCX", """
from docx import Document
doc = Document()
doc.add_heading("Contrat de prestation", level=1)
doc.add_paragraph("Le Client sengage a respecter les conditions de paiement sous 30 jours.")
doc.add_paragraph("Le Client peut resilier le contrat avec un preavis de 30 jours.")
doc.add_paragraph("Le prestataire sengage a livrer dans les delais convenus.")
doc.save("/mnt/data/contrat.docx")
import os; print("Created:", os.path.getsize("/mnt/data/contrat.docx"), "bytes")
""", session_id=sid)

if ok:
    sid, ok = exec_code("DOCX-2: Unpack", """
import subprocess
r = subprocess.run(["python3", "/opt/skills/docx/scripts/office/unpack.py",
    "/mnt/data/contrat.docx", "/mnt/data/contrat_unpacked"],
    capture_output=True, text=True)
print("rc:", r.returncode)
print(r.stdout.strip()[:200])
assert r.returncode == 0, r.stderr[:300]
import os; assert os.path.isdir("/mnt/data/contrat_unpacked/word")
print("Unpacked OK")
""", session_id=sid)

if ok:
    sid, ok = exec_code("DOCX-3: Tracked replace (2 pairs)", """
import subprocess
r = subprocess.run(["python3", "/opt/skills/docx/scripts/tracked_replace.py",
    "/mnt/data/contrat_unpacked",
    "--old", "Le Client", "--new", "L Utilisateur",
    "--old", "30 jours", "--new", "15 jours ouvres",
    "--author", "AI-Agent"],
    capture_output=True, text=True)
print(r.stdout.strip())
assert r.returncode == 0, r.stderr[:300]
""", session_id=sid)

if ok:
    sid, ok = exec_code("DOCX-4: Pack + validate", """
import subprocess
r = subprocess.run(["python3", "/opt/skills/docx/scripts/office/pack.py",
    "/mnt/data/contrat_unpacked", "-o", "/mnt/data/contrat_redline.docx"],
    capture_output=True, text=True)
print("pack rc:", r.returncode, r.stdout.strip()[:100])
assert r.returncode == 0, r.stderr[:300]
r2 = subprocess.run(["python3", "/opt/skills/docx/scripts/office/validate.py",
    "/mnt/data/contrat_redline.docx"], capture_output=True, text=True)
print("validate rc:", r2.returncode, r2.stdout.strip()[:100])
import os; print("Redlined DOCX:", os.path.getsize("/mnt/data/contrat_redline.docx"), "bytes")
""", session_id=sid)

if ok:
    sid, ok = exec_code("DOCX-5: Accept tracked changes (soffice)", """
import subprocess, os
r = subprocess.run(["python3", "/opt/skills/docx/scripts/accept_changes.py",
    "/mnt/data/contrat_redline.docx", "/mnt/data/contrat_clean.docx"],
    capture_output=True, text=True, timeout=120)
print("rc:", r.returncode)
print("stdout:", r.stdout.strip()[:200])
if r.stderr: print("stderr:", r.stderr.strip()[:200])
if os.path.exists("/mnt/data/contrat_clean.docx"):
    print("Clean DOCX:", os.path.getsize("/mnt/data/contrat_clean.docx"), "bytes")
else:
    print("File not created - checking accept_changes output")
""", session_id=sid, timeout=180)

sid, _ = exec_code("DOCX-6: Convert to PDF (soffice)", """
import sys, subprocess, os
sys.path.insert(0, "/opt/skills/docx/scripts")
from office.soffice import run_soffice
src = "/mnt/data/contrat_redline.docx"
if not os.path.exists(src):
    src = "/mnt/data/contrat.docx"
r = run_soffice(["--headless", "--convert-to", "pdf", "--outdir", "/mnt/data", src],
    capture_output=True, text=True, timeout=120)
print("soffice:", r.stdout.strip()[:200])
pdfs = [f for f in os.listdir("/mnt/data") if f.endswith(".pdf")]
print("PDFs in /mnt/data:", pdfs)
for p in pdfs:
    print(f"  {p}: {os.path.getsize('/mnt/data/' + p)} bytes")
""", session_id=sid, timeout=180)

# ====================================================================
print("\n=== Pipeline XLSX: create -> recalc -> PDF ===")
# ====================================================================
sid2 = str(uuid.uuid4())

sid2, ok = exec_code("XLSX-1: Create with formulas", """
from openpyxl import Workbook
from openpyxl.styles import Font
wb = Workbook()
ws = wb.active; ws.title = "Budget"
ws.append(["Poste", "Montant"])
for p, v in [("Salaires",45000),("Loyer",8000),("Marketing",12000),("IT",5000)]:
    ws.append([p, v])
ws.cell(row=6, column=1, value="TOTAL").font = Font(bold=True)
ws.cell(row=6, column=2).value = "=SUM(B2:B5)"
ws.cell(row=6, column=2).number_format = "#,##0"
wb.save("/mnt/data/budget.xlsx")
import os; print("Created:", os.path.getsize("/mnt/data/budget.xlsx"), "bytes")
""", session_id=sid2)

if ok:
    sid2, ok = exec_code("XLSX-2: Recalc formulas (soffice)", """
import subprocess
r = subprocess.run(["python3", "/opt/skills/xlsx/scripts/recalc.py", "/mnt/data/budget.xlsx"],
    capture_output=True, text=True, timeout=120)
print("rc:", r.returncode)
print(r.stdout.strip()[:300])
if r.stderr: print("stderr:", r.stderr.strip()[:200])
""", session_id=sid2, timeout=180)

if ok:
    sid2, _ = exec_code("XLSX-3: Export PDF (soffice)", """
import sys, subprocess, os
sys.path.insert(0, "/opt/skills/xlsx/scripts")
from office.soffice import run_soffice
r = run_soffice(["--headless", "--convert-to", "pdf", "--outdir", "/mnt/data", "/mnt/data/budget.xlsx"],
    capture_output=True, text=True, timeout=120)
print("soffice:", r.stdout.strip()[:200])
if os.path.exists("/mnt/data/budget.pdf"):
    print("PDF:", os.path.getsize("/mnt/data/budget.pdf"), "bytes")
else:
    print("PDFs:", [f for f in os.listdir("/mnt/data") if f.endswith(".pdf")])
""", session_id=sid2, timeout=180)

# ====================================================================
print("\n=== Pipeline PDF: create -> extract text -> extract tables -> merge -> split -> watermark ===")
# ====================================================================
sid3 = str(uuid.uuid4())

sid3, ok = exec_code("PDF-1: Create test PDFs", """
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
import os

# PDF with text
c = canvas.Canvas("/mnt/data/doc1.pdf", pagesize=A4)
c.setFont("Helvetica", 16); c.drawString(100, 750, "Document partie 1")
c.setFont("Helvetica", 12); c.drawString(100, 700, "Ceci est le premier document de test.")
c.drawString(100, 680, "Il contient du texte sur plusieurs lignes.")
c.save()

# PDF with table
doc = SimpleDocTemplate("/mnt/data/doc2.pdf", pagesize=A4)
data = [["Produit","Qty","Prix"],["Widget A","10","25.50"],["Widget B","5","42.00"],["Widget C","20","15.75"]]
t = Table(data)
t.setStyle(TableStyle([("GRID",(0,0),(-1,-1),1,colors.black),("BACKGROUND",(0,0),(-1,0),colors.lightblue)]))
doc.build([t])

print("doc1.pdf:", os.path.getsize("/mnt/data/doc1.pdf"), "bytes")
print("doc2.pdf:", os.path.getsize("/mnt/data/doc2.pdf"), "bytes")
""", session_id=sid3)

if ok:
    sid3, ok = exec_code("PDF-2: Extract text (pdfplumber)", """
import pdfplumber
with pdfplumber.open("/mnt/data/doc1.pdf") as pdf:
    text = pdf.pages[0].extract_text()
print("Extracted text:")
print(text[:300])
assert "Document partie 1" in text
""", session_id=sid3)

if ok:
    sid3, ok = exec_code("PDF-3: Extract tables (pdfplumber)", """
import pdfplumber, pandas as pd
with pdfplumber.open("/mnt/data/doc2.pdf") as pdf:
    tables = pdf.pages[0].extract_tables()
print("Tables found:", len(tables))
if tables:
    df = pd.DataFrame(tables[0][1:], columns=tables[0][0])
    print(df.to_string())
    df.to_excel("/mnt/data/extracted_table.xlsx", index=False)
    import os; print("Excel:", os.path.getsize("/mnt/data/extracted_table.xlsx"), "bytes")
""", session_id=sid3)

if ok:
    sid3, ok = exec_code("PDF-4: Merge (pypdf)", """
from pypdf import PdfReader, PdfWriter
import os
writer = PdfWriter()
for f in ["doc1.pdf", "doc2.pdf"]:
    reader = PdfReader("/mnt/data/" + f)
    for page in reader.pages:
        writer.add_page(page)
with open("/mnt/data/merged.pdf", "wb") as out:
    writer.write(out)
r = PdfReader("/mnt/data/merged.pdf")
print("Merged:", len(r.pages), "pages,", os.path.getsize("/mnt/data/merged.pdf"), "bytes")
""", session_id=sid3)

if ok:
    sid3, ok = exec_code("PDF-5: Split page 2", """
from pypdf import PdfReader, PdfWriter
import os
reader = PdfReader("/mnt/data/merged.pdf")
writer = PdfWriter()
writer.add_page(reader.pages[1])
with open("/mnt/data/page2.pdf", "wb") as f:
    writer.write(f)
print("Page 2 extracted:", os.path.getsize("/mnt/data/page2.pdf"), "bytes")
""", session_id=sid3)

if ok:
    sid3, ok = exec_code("PDF-6: Watermark", """
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.pagesizes import A4
from pypdf import PdfReader, PdfWriter
import io, os
packet = io.BytesIO()
c = rl_canvas.Canvas(packet, pagesize=A4)
c.saveState(); c.setFillAlpha(0.3); c.setFillGray(0.5)
c.setFont("Helvetica", 60); c.translate(300, 400); c.rotate(45)
c.drawCentredString(0, 0, "BROUILLON"); c.restoreState(); c.save()
packet.seek(0)
wm = PdfReader(packet).pages[0]
reader = PdfReader("/mnt/data/merged.pdf")
writer = PdfWriter()
for page in reader.pages:
    page.merge_page(wm); writer.add_page(page)
with open("/mnt/data/watermarked.pdf", "wb") as f:
    writer.write(f)
print("Watermarked:", os.path.getsize("/mnt/data/watermarked.pdf"), "bytes")
""", session_id=sid3)

if ok:
    sid3, _ = exec_code("PDF-7: qpdf linearize", """
import subprocess, os
r = subprocess.run(["qpdf", "--linearize", "/mnt/data/watermarked.pdf", "/mnt/data/optimized.pdf"],
    capture_output=True, text=True)
print("rc:", r.returncode)
s1 = os.path.getsize("/mnt/data/watermarked.pdf")
s2 = os.path.getsize("/mnt/data/optimized.pdf")
print("Before:", s1, "After:", s2)
""", session_id=sid3)

# ====================================================================
print("\n=== Pipeline FFmpeg: generate -> convert -> extract -> analyse ===")
# ====================================================================
sid4 = str(uuid.uuid4())

sid4, ok = exec_code("MEDIA-1: Generate WAV + MP3", """
import subprocess, os
subprocess.run(["ffmpeg","-y","-f","lavfi","-i","sine=frequency=440:duration=3",
    "-f","wav","/mnt/data/tone.wav"], check=True, capture_output=True)
subprocess.run(["ffmpeg","-y","-i","/mnt/data/tone.wav","-c:a","libmp3lame","-b:a","192k",
    "/mnt/data/tone.mp3"], check=True, capture_output=True)
print("WAV:", os.path.getsize("/mnt/data/tone.wav"), "bytes")
print("MP3:", os.path.getsize("/mnt/data/tone.mp3"), "bytes")
""", session_id=sid4)

if ok:
    sid4, ok = exec_code("MEDIA-2: Create video + extract frame", """
import subprocess, os
subprocess.run(["ffmpeg","-y","-f","lavfi","-i","color=c=blue:s=640x480:d=3",
    "-c:v","libx264","-t","3","/mnt/data/video.mp4"], check=True, capture_output=True)
print("Video:", os.path.getsize("/mnt/data/video.mp4"), "bytes")
subprocess.run(["ffmpeg","-y","-i","/mnt/data/video.mp4","-ss","00:00:01",
    "-frames:v","1","/mnt/data/frame.png"], check=True, capture_output=True)
print("Frame:", os.path.getsize("/mnt/data/frame.png"), "bytes")
""", session_id=sid4)

if ok:
    sid4, ok = exec_code("MEDIA-3: ffprobe analysis", """
import subprocess, json
r = subprocess.run(["ffprobe","-v","quiet","-print_format","json",
    "-show_format","-show_streams","/mnt/data/tone.mp3"],
    capture_output=True, text=True, check=True)
info = json.loads(r.stdout)
print("Format:", info["format"]["format_name"])
print("Duration:", info["format"]["duration"], "s")
print("Codec:", info["streams"][0]["codec_name"])
""", session_id=sid4)

if ok:
    sid4, _ = exec_code("MEDIA-4: Pillow image pipeline", """
from PIL import Image, ImageDraw
import os
img = Image.new("RGB", (800, 600))
draw = ImageDraw.Draw(img)
for y in range(600):
    draw.line([(0,y),(799,y)], fill=(int(50*y/600),int(100+100*y/600),int(200+55*y/600)))
draw.text((320, 280), "Test Agent", fill=(0,0,0))
img.save("/mnt/data/gradient.png")
img.thumbnail((400, 300))
img.convert("RGB").save("/mnt/data/thumb.jpg", quality=85)
print("PNG:", os.path.getsize("/mnt/data/gradient.png"), "bytes")
print("JPEG thumb:", os.path.getsize("/mnt/data/thumb.jpg"), "bytes")
""", session_id=sid4)

# ====================================================================
print("\n=== Pipeline DataViz: generate -> analyse -> viz -> export ===")
# ====================================================================
sid5 = str(uuid.uuid4())

sid5, ok = exec_code("DATAVIZ-1: Generate + analyse", """
import pandas as pd, numpy as np
np.random.seed(42)
n = 500
df = pd.DataFrame({
    "date": pd.date_range("2026-01-01", periods=n),
    "cat": np.random.choice(["A","B","C"], n),
    "montant": np.random.normal(1000, 300, n).round(2),
    "qty": np.random.randint(1, 50, n)
})
df.to_csv("/mnt/data/dataset.csv", index=False)
print("Shape:", df.shape)
print("Mean montant:", round(df.montant.mean(), 2))
print("Categories:", df.cat.value_counts().to_dict())
""", session_id=sid5)

if ok:
    sid5, ok = exec_code("DATAVIZ-2: Dashboard 4 charts", """
import pandas as pd, numpy as np
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns
from scipy.stats import linregress
import os
df = pd.read_csv("/mnt/data/dataset.csv", parse_dates=["date"])
df["mois"] = df["date"].dt.month
fig, axes = plt.subplots(2, 2, figsize=(14, 10))
monthly = df.groupby("mois")["montant"].sum()
axes[0,0].plot(monthly.index, monthly.values, marker="o"); axes[0,0].set_title("CA mensuel")
cat_sum = df.groupby("cat")["montant"].sum()
axes[0,1].pie(cat_sum, labels=cat_sum.index, autopct="%1.1f%%"); axes[0,1].set_title("Par categorie")
sns.boxplot(data=df, x="cat", y="montant", ax=axes[1,0]); axes[1,0].set_title("Montants")
axes[1,1].scatter(df["qty"], df["montant"], alpha=0.3, s=10)
sl, ic, r, p, se = linregress(df["qty"], df["montant"])
axes[1,1].plot([0,50], [ic, sl*50+ic], "r-"); axes[1,1].set_title("Qty vs Montant")
plt.tight_layout()
plt.savefig("/mnt/data/dashboard.png", dpi=150, bbox_inches="tight"); plt.close()
print("Dashboard:", os.path.getsize("/mnt/data/dashboard.png"), "bytes")
""", session_id=sid5)

if ok:
    sid5, ok = exec_code("DATAVIZ-3: ANOVA + t-test", """
import pandas as pd, numpy as np
from scipy import stats
df = pd.read_csv("/mnt/data/dataset.csv")
groups = [g["montant"].values for _, g in df.groupby("cat")]
f, p = stats.f_oneway(*groups)
print("ANOVA: F=%.4f p=%.6f Sig=%s" % (f, p, "Oui" if p < 0.05 else "Non"))
np.random.seed(42)
avant = np.random.normal(72, 12, 50); apres = np.random.normal(78, 11, 50)
t, p2 = stats.ttest_rel(avant, apres)
print("t-test: t=%.4f p=%.6f Sig=%s" % (t, p2, "Oui" if p2 < 0.05 else "Non"))
""", session_id=sid5)

if ok:
    sid5, _ = exec_code("DATAVIZ-4: Multi-sheet Excel export", """
import pandas as pd, os
df = pd.read_csv("/mnt/data/dataset.csv", parse_dates=["date"])
df["mois"] = df["date"].dt.month
with pd.ExcelWriter("/mnt/data/rapport.xlsx", engine="openpyxl") as w:
    df.to_excel(w, sheet_name="Brut", index=False)
    pivot = df.pivot_table(values="montant", index="mois", columns="cat", aggfunc="sum")
    pivot.to_excel(w, sheet_name="Pivot")
    st = df.groupby("cat")["montant"].describe()
    st.to_excel(w, sheet_name="Stats")
print("Rapport:", os.path.getsize("/mnt/data/rapport.xlsx"), "bytes")
""", session_id=sid5)

# ====================================================================
print("\n" + "=" * 60)
print("FINAL RESULTS")
print("=" * 60)
print(f"Passed: {PASS}")
print(f"Failed: {FAIL}")
print(f"Total:  {PASS + FAIL}")
print("=" * 60)
for status, label, err in RESULTS:
    mark = "\u2713" if status == "PASS" else "\u2717"
    line = f"  {mark} {label}"
    if err:
        line += f"  -- {err}"
    print(line)
print()
sys.exit(0 if FAIL == 0 else 1)
